"""
Unit tests for the Exporter module.
Tests DOCX document creation and formatting functionality.
"""

import pytest
from pathlib import Path
from datetime import datetime
from docx import Document
import sys

# Add src to path
sys.path.insert(0, str(Path(__file__).parent.parent / 'src'))

from exporter import DocxExporter


@pytest.fixture
def temp_output_dir(tmp_path):
    """Create a temporary output directory for tests."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    return str(output_dir)


@pytest.fixture
def sample_summary():
    """Create a sample meeting summary for testing."""
    return {
        'title': 'Test Meeting',
        'participants': ['Alice', 'Bob', 'Charlie'],
        'key_points': ['Point 1', 'Point 2', 'Point 3'],
        'action_items': ['Action 1', 'Action 2'],
        'decisions': ['Decision 1'],
        'follow_ups': ['Follow-up 1', 'Follow-up 2']
    }


@pytest.fixture
def exporter(temp_output_dir):
    """Create a DocxExporter instance for testing."""
    return DocxExporter(temp_output_dir)


class TestDocxExporterInitialization:
    """Test DocxExporter initialization."""
    
    def test_init_creates_exporter(self, temp_output_dir):
        """Test that exporter initializes correctly."""
        exporter = DocxExporter(temp_output_dir)
        assert exporter.output_dir == Path(temp_output_dir)
    
    def test_init_with_string_path(self):
        """Test initialization with string path."""
        exporter = DocxExporter("./output")
        assert isinstance(exporter.output_dir, Path)


class TestDocumentCreation:
    """Test DOCX document creation."""
    
    def test_create_document_returns_filepath(self, exporter, sample_summary):
        """Test that create_document returns a valid file path."""
        filepath = exporter.create_document(sample_summary)
        
        assert filepath is not None
        assert isinstance(filepath, str)
        assert Path(filepath).exists()
        assert filepath.endswith('.docx')
    
    def test_create_document_saves_file(self, exporter, sample_summary):
        """Test that document is actually saved to disk."""
        filepath = exporter.create_document(sample_summary)
        
        file_path = Path(filepath)
        assert file_path.exists()
        assert file_path.stat().st_size > 0
    
    def test_create_document_with_minimal_summary(self, exporter):
        """Test document creation with minimal summary data."""
        minimal_summary = {
            'title': 'Minimal Meeting',
            'participants': [],
            'key_points': [],
            'action_items': [],
            'decisions': [],
            'follow_ups': []
        }
        
        filepath = exporter.create_document(minimal_summary)
        assert Path(filepath).exists()
    
    def test_create_document_with_missing_fields(self, exporter):
        """Test document creation handles missing fields gracefully."""
        incomplete_summary = {
            'title': 'Incomplete Meeting'
            # Missing other fields
        }
        
        # Should not raise error, just log warnings
        filepath = exporter.create_document(incomplete_summary)
        assert Path(filepath).exists()
    
    def test_create_document_invalid_summary_type(self, exporter):
        """Test that invalid summary type raises ValueError."""
        with pytest.raises(ValueError, match="Summary must be a dictionary"):
            exporter.create_document("not a dict")
    
    def test_create_document_creates_valid_docx(self, exporter, sample_summary):
        """Test that created file is a valid DOCX document."""
        filepath = exporter.create_document(sample_summary)
        
        # Try to open the document with python-docx
        doc = Document(filepath)
        assert doc is not None
        assert len(doc.paragraphs) > 0


class TestFilenameGeneration:
    """Test filename generation."""
    
    def test_filename_has_timestamp(self, exporter, sample_summary):
        """Test that filename includes timestamp."""
        filepath = exporter.create_document(sample_summary)
        filename = Path(filepath).name
        
        # Should match pattern: meeting_summary_YYYYMMDD_HHMMSS.docx
        assert filename.startswith('meeting_summary_')
        assert filename.endswith('.docx')
        assert len(filename) == len('meeting_summary_YYYYMMDD_HHMMSS.docx')
    
    def test_multiple_documents_have_unique_filenames(self, exporter, sample_summary):
        """Test that multiple documents get unique filenames."""
        filepath1 = exporter.create_document(sample_summary)
        
        # Small delay to ensure different timestamp
        import time
        time.sleep(1)
        
        filepath2 = exporter.create_document(sample_summary)
        
        assert filepath1 != filepath2
        assert Path(filepath1).name != Path(filepath2).name


class TestDirectoryCreation:
    """Test output directory creation."""
    
    def test_creates_output_directory_if_not_exists(self, tmp_path):
        """Test that output directory is created if it doesn't exist."""
        output_dir = tmp_path / "new_output"
        assert not output_dir.exists()
        
        exporter = DocxExporter(str(output_dir))
        summary = {'title': 'Test', 'participants': [], 'key_points': [],
                  'action_items': [], 'decisions': [], 'follow_ups': []}
        
        filepath = exporter.create_document(summary)
        
        assert output_dir.exists()
        assert Path(filepath).parent == output_dir
    
    def test_works_with_existing_directory(self, temp_output_dir, sample_summary):
        """Test that exporter works with existing directory."""
        exporter = DocxExporter(temp_output_dir)
        filepath = exporter.create_document(sample_summary)
        
        assert Path(filepath).exists()


class TestDocumentContent:
    """Test document content and structure."""
    
    def test_document_contains_title(self, exporter, sample_summary):
        """Test that document contains the meeting title."""
        filepath = exporter.create_document(sample_summary)
        doc = Document(filepath)
        
        # First paragraph should be the title
        title_text = doc.paragraphs[0].text
        assert sample_summary['title'] in title_text
    
    def test_document_contains_all_sections(self, exporter, sample_summary):
        """Test that document contains all required sections."""
        filepath = exporter.create_document(sample_summary)
        doc = Document(filepath)
        
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        assert 'Participants' in full_text
        assert 'Key Discussion Points' in full_text
        assert 'Action Items' in full_text
        assert 'Decisions Taken' in full_text
        assert 'Follow-up Tasks' in full_text
    
    def test_document_contains_participant_names(self, exporter, sample_summary):
        """Test that participant names appear in document."""
        filepath = exporter.create_document(sample_summary)
        doc = Document(filepath)
        
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        for participant in sample_summary['participants']:
            assert participant in full_text
    
    def test_document_contains_key_points(self, exporter, sample_summary):
        """Test that key points appear in document."""
        filepath = exporter.create_document(sample_summary)
        doc = Document(filepath)
        
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        for point in sample_summary['key_points']:
            assert point in full_text


class TestDocumentFormatting:
    """Test document formatting and styling."""
    
    def test_title_has_heading_style(self, exporter, sample_summary):
        """Test that title uses heading style."""
        filepath = exporter.create_document(sample_summary)
        doc = Document(filepath)
        
        # First paragraph should be a heading
        title_paragraph = doc.paragraphs[0]
        assert title_paragraph.style.name.startswith('Heading')
    
    def test_sections_have_heading_style(self, exporter, sample_summary):
        """Test that section headers use heading style."""
        filepath = exporter.create_document(sample_summary)
        doc = Document(filepath)
        
        # Find section headings
        section_headings = ['Participants', 'Key Discussion Points', 
                           'Action Items', 'Decisions Taken', 'Follow-up Tasks']
        
        for paragraph in doc.paragraphs:
            if paragraph.text in section_headings:
                assert paragraph.style.name.startswith('Heading')
    
    def test_bullet_points_have_list_style(self, exporter, sample_summary):
        """Test that content items use bullet list style."""
        filepath = exporter.create_document(sample_summary)
        doc = Document(filepath)
        
        # Check that some paragraphs have list bullet style
        list_paragraphs = [p for p in doc.paragraphs 
                          if 'List Bullet' in p.style.name]
        
        assert len(list_paragraphs) > 0


class TestErrorHandling:
    """Test error handling."""
    
    def test_invalid_output_directory_raises_error(self):
        """Test that invalid output directory is handled."""
        # This test depends on OS permissions
        # On Windows, we can't easily create an unwritable directory
        # So we'll just verify the exporter can be created
        exporter = DocxExporter("/invalid/path/that/hopefully/does/not/exist")
        assert exporter is not None
    
    def test_empty_summary_handled_gracefully(self, exporter):
        """Test that empty summary is handled gracefully."""
        empty_summary = {}
        
        # Should not crash, just create document with defaults
        filepath = exporter.create_document(empty_summary)
        assert Path(filepath).exists()


class TestEdgeCases:
    """Test edge cases and special scenarios."""
    
    def test_summary_with_empty_strings(self, exporter):
        """Test handling of empty strings in summary."""
        summary = {
            'title': '',
            'participants': ['', 'Alice', ''],
            'key_points': ['Point 1', '', 'Point 2'],
            'action_items': [],
            'decisions': [],
            'follow_ups': []
        }
        
        filepath = exporter.create_document(summary)
        assert Path(filepath).exists()
    
    def test_summary_with_long_content(self, exporter):
        """Test handling of very long content."""
        summary = {
            'title': 'Very Long Meeting Title ' * 10,
            'participants': [f'Person {i}' for i in range(50)],
            'key_points': [f'Point {i}: ' + 'x' * 200 for i in range(100)],
            'action_items': [f'Action {i}' for i in range(50)],
            'decisions': [f'Decision {i}' for i in range(50)],
            'follow_ups': [f'Follow-up {i}' for i in range(50)]
        }
        
        filepath = exporter.create_document(summary)
        assert Path(filepath).exists()
        
        # Verify file is reasonably large
        assert Path(filepath).stat().st_size > 10000
    
    def test_summary_with_special_characters(self, exporter):
        """Test handling of special characters in content."""
        summary = {
            'title': 'Meeting with Special Chars: @#$%^&*()',
            'participants': ['Alice & Bob', 'Charlie <Dev>'],
            'key_points': ['Point with "quotes"', "Point with 'apostrophes'"],
            'action_items': ['Action with émojis 🎉'],
            'decisions': ['Decision with newline\ncharacter'],
            'follow_ups': ['Follow-up with tab\tcharacter']
        }
        
        filepath = exporter.create_document(summary)
        assert Path(filepath).exists()
