"""
Exporter module for MeetDocs AI automation system.
Creates formatted DOCX documents from meeting summaries.
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


logger = logging.getLogger(__name__)


class DocxExporter:
    """
    Manages DOCX document generation for meeting summaries.
    
    Creates professionally formatted documents with proper styling,
    including headings, bullet points, and consistent formatting.
    """
    
    def __init__(self, output_dir: str):
        """
        Initialize the DocxExporter with output directory.
        
        Args:
            output_dir: Directory path where DOCX files will be saved
        """
        self.output_dir = Path(output_dir)
        logger.info(f"DocxExporter initialized with output directory: {self.output_dir}")
    
    def create_document(self, summary: Dict[str, any]) -> str:
        """
        Generate a formatted DOCX document from meeting summary.
        
        Creates a professional document with proper styling and structure,
        saves it with a timestamped filename, and returns the file path.
        
        Args:
            summary: Dictionary containing structured summary with keys:
                    - title: Meeting title
                    - participants: List of participant names
                    - key_points: List of key discussion points
                    - action_items: List of action items
                    - decisions: List of decisions taken
                    - follow_ups: List of follow-up tasks
        
        Returns:
            Absolute path to the created DOCX file
            
        Raises:
            ValueError: If summary structure is invalid
            IOError: If file cannot be written
        """
        logger.info("Creating DOCX document from summary")
        
        # Validate summary structure
        self._validate_summary(summary)
        
        # Ensure output directory exists
        self._ensure_output_directory()
        
        # Create new document
        doc = Document()
        
        # Apply formatting
        self._apply_formatting(doc)
        
        # Add title
        self._add_title(doc, summary.get('title', 'Meeting Summary'))
        
        # Add sections
        self._add_section(doc, 'Participants', summary.get('participants', []))
        self._add_section(doc, 'Key Discussion Points', summary.get('key_points', []))
        self._add_section(doc, 'Action Items', summary.get('action_items', []))
        self._add_section(doc, 'Decisions Taken', summary.get('decisions', []))
        self._add_section(doc, 'Follow-up Tasks', summary.get('follow_ups', []))
        
        # Generate filename and save
        filename = self._generate_filename()
        filepath = self.output_dir / filename
        
        try:
            doc.save(str(filepath))
            logger.info(f"Document saved successfully: {filepath}")
            return str(filepath.absolute())
            
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            raise IOError(f"Could not save document to {filepath}: {e}")

    
    def _validate_summary(self, summary: Dict[str, any]):
        """
        Validate that summary has required structure.
        
        Args:
            summary: Summary dictionary to validate
            
        Raises:
            ValueError: If summary is invalid
        """
        if not isinstance(summary, dict):
            raise ValueError("Summary must be a dictionary")
        
        required_fields = ['title', 'participants', 'key_points', 
                          'action_items', 'decisions', 'follow_ups']
        
        for field in required_fields:
            if field not in summary:
                logger.warning(f"Missing field in summary: {field}")
                # Don't raise error, just log warning
        
        logger.debug("Summary structure validated")
    
    def _ensure_output_directory(self):
        """
        Create output directory if it doesn't exist.
        
        Raises:
            IOError: If directory cannot be created
        """
        try:
            self.output_dir.mkdir(parents=True, exist_ok=True)
            logger.debug(f"Output directory ensured: {self.output_dir}")
            
        except Exception as e:
            logger.error(f"Failed to create output directory: {e}")
            raise IOError(f"Could not create directory {self.output_dir}: {e}")
    
    def _apply_formatting(self, doc: Document):
        """
        Apply professional formatting to the document.
        
        Sets up document-wide formatting including margins and default styles.
        
        Args:
            doc: Document object to format
        """
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        logger.debug("Document formatting applied")
    
    def _add_title(self, doc: Document, title: str):
        """
        Add formatted title to the document.
        
        Applies Heading 1 style with bold formatting and 16pt font.
        
        Args:
            doc: Document object
            title: Meeting title text
        """
        # Use default title if empty
        if not title or not title.strip():
            title = 'Meeting Summary'
        
        heading = doc.add_heading(title, level=1)
        
        # Apply Heading 1 formatting
        if heading.runs:
            heading_format = heading.runs[0]
            heading_format.font.size = Pt(16)
            heading_format.font.bold = True
            heading_format.font.name = 'Calibri'
        
        logger.debug(f"Title added: {title}")
    
    def _add_section(self, doc: Document, heading: str, content: List[str]):
        """
        Add a section with heading and bullet points.
        
        Creates a section header with Heading 2 style and adds content
        as bullet points. Handles empty content gracefully.
        
        Args:
            doc: Document object
            heading: Section heading text
            content: List of items to add as bullet points
        """
        # Add section heading
        section_heading = doc.add_heading(heading, level=2)
        
        # Apply Heading 2 formatting
        heading_format = section_heading.runs[0]
        heading_format.font.size = Pt(14)
        heading_format.font.bold = True
        heading_format.font.name = 'Calibri'
        
        # Add content as bullet points
        if content and len(content) > 0:
            for item in content:
                if item and str(item).strip():  # Skip empty items
                    paragraph = doc.add_paragraph(str(item), style='List Bullet')
                    
                    # Set paragraph formatting
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_after = Pt(6)
                    paragraph_format.line_spacing = 1.15
                    paragraph_format.left_indent = Inches(0.5)
                    
                    # Set font for bullet text
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
        else:
            # Add placeholder if no content
            paragraph = doc.add_paragraph('No items recorded', style='Normal')
            paragraph.runs[0].font.italic = True
            paragraph.runs[0].font.name = 'Calibri'
            paragraph.runs[0].font.size = Pt(11)
        
        logger.debug(f"Section added: {heading} with {len(content) if content else 0} items")
    
    def _generate_filename(self) -> str:
        """
        Generate timestamped filename for the document.
        
        Creates a unique filename using the current timestamp to prevent
        overwriting previous summaries.
        
        Returns:
            Filename string in format: meeting_summary_YYYYMMDD_HHMMSS.docx
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'meeting_summary_{timestamp}.docx'
        
        logger.debug(f"Generated filename: {filename}")
        return filename
